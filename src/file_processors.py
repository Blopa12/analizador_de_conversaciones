"""
File processing modules for different file formats (xlsx, docx, pdf)
"""

import pandas as pd
import docx
import PyPDF2
from typing import Dict, List, Any, Optional
import io
from pathlib import Path


class FileProcessor:
    """Base class for file processing"""

    def __init__(self):
        self.supported_formats = ['.xlsx', '.xls', '.csv', '.docx', '.pdf']

    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a file and return structured data

        Args:
            file_path: Path to the file to process

        Returns:
            Dictionary containing:
            - file_type: Type of file processed
            - raw_data: Raw extracted data
            - metadata: File metadata
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        if file_extension in ['.xlsx', '.xls']:
            return self._process_excel(file_path)
        elif file_extension == '.csv':
            return self._process_csv(file_path)
        elif file_extension == '.docx':
            return self._process_docx(file_path)
        elif file_extension == '.pdf':
            return self._process_pdf(file_path)

    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel file (xlsx or xls)"""
        try:
            # Try to read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)

            result = {
                'file_type': 'excel',
                'raw_data': excel_data,
                'metadata': {
                    'sheets': list(excel_data.keys()),
                    'total_sheets': len(excel_data),
                    'file_size': file_path.stat().st_size,
                    'file_name': file_path.name
                }
            }

            # Add row counts for each sheet
            for sheet_name, df in excel_data.items():
                result['metadata'][f'{sheet_name}_rows'] = len(df)
                result['metadata'][f'{sheet_name}_columns'] = len(df.columns)

            return result

        except Exception as e:
            raise Exception(f"Error processing Excel file: {str(e)}")

    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            # Try to read CSV with different encodings and separators
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            separators = [',', ';', '\t', '|']

            df = None
            encoding_used = None
            separator_used = None

            for encoding in encodings:
                for separator in separators:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=separator)
                        # Check if we got meaningful data (more than 1 column and reasonable rows)
                        if len(df.columns) > 1 and len(df) > 0:
                            encoding_used = encoding
                            separator_used = separator
                            break
                    except:
                        continue
                if df is not None:
                    break

            if df is None:
                # Fallback: try with default parameters
                df = pd.read_csv(file_path)
                encoding_used = 'default'
                separator_used = ','

            # Convert to same format as Excel processing
            csv_data = {'Sheet1': df}  # Treat CSV as single sheet

            result = {
                'file_type': 'csv',
                'raw_data': csv_data,
                'metadata': {
                    'sheets': ['Sheet1'],
                    'total_sheets': 1,
                    'encoding': encoding_used,
                    'separator': separator_used,
                    'file_size': file_path.stat().st_size,
                    'file_name': file_path.name,
                    'Sheet1_rows': len(df),
                    'Sheet1_columns': len(df.columns)
                }
            }

            return result

        except Exception as e:
            raise Exception(f"Error processing CSV file: {str(e)}")

    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process Word document"""
        try:
            doc = docx.Document(file_path)

            paragraphs = []
            tables = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                'file_type': 'docx',
                'raw_data': {
                    'paragraphs': paragraphs,
                    'tables': tables
                },
                'metadata': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables),
                    'file_size': file_path.stat().st_size,
                    'file_name': file_path.name
                }
            }

        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")

    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append({
                                'page': page_num + 1,
                                'content': text.strip()
                            })
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {str(e)}")

                return {
                    'file_type': 'pdf',
                    'raw_data': text_content,
                    'metadata': {
                        'page_count': len(pdf_reader.pages),
                        'pages_with_text': len(text_content),
                        'file_size': file_path.stat().st_size,
                        'file_name': file_path.name
                    }
                }

        except Exception as e:
            raise Exception(f"Error processing PDF file: {str(e)}")


class DataStructureAnalyzer:
    """Analyzes the structure of extracted data to identify patterns"""

    def analyze_structure(self, processed_file: Dict[str, Any]) -> Dict[str, Any]:
        """
        Analyze the structure of processed file data

        Args:
            processed_file: Output from FileProcessor.process_file()

        Returns:
            Dictionary with structure analysis results
        """
        file_type = processed_file['file_type']

        if file_type in ['excel', 'csv']:
            return self._analyze_excel_structure(processed_file)
        elif file_type == 'docx':
            return self._analyze_docx_structure(processed_file)
        elif file_type == 'pdf':
            return self._analyze_pdf_structure(processed_file)

    def _analyze_excel_structure(self, processed_file: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze Excel/CSV file structure"""
        raw_data = processed_file['raw_data']
        analysis = {
            'structure_type': 'tabular',
            'sheets': {},
            'potential_identifiers': [],
            'potential_content_columns': []
        }

        for sheet_name, df in raw_data.items():
            if df.empty:
                continue

            sheet_analysis = {
                'columns': list(df.columns),
                'sample_data': df.head(3).to_dict('records') if len(df) > 0 else [],
                'data_types': df.dtypes.to_dict(),
                'potential_id_columns': [],
                'potential_content_columns': []
            }

            # Identify potential ID columns
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['id', 'ticket', 'conversation', 'numero', 'folio']):
                    sheet_analysis['potential_id_columns'].append(col)

            # Identify potential content columns
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['content', 'message', 'description', 'comment', 'text', 'contenido', 'mensaje', 'descripcion']):
                    sheet_analysis['potential_content_columns'].append(col)

            analysis['sheets'][sheet_name] = sheet_analysis

        return analysis

    def _analyze_docx_structure(self, processed_file: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze Word document structure"""
        raw_data = processed_file['raw_data']

        # Look for patterns in paragraphs
        patterns = {
            'ticket_patterns': [],
            'conversation_patterns': [],
            'structured_sections': []
        }

        for i, paragraph in enumerate(raw_data['paragraphs']):
            para_lower = paragraph.lower()

            # Look for ticket-like patterns
            if any(keyword in para_lower for keyword in ['ticket', 'caso', 'incidente', '#']):
                patterns['ticket_patterns'].append({'line': i, 'content': paragraph})

            # Look for conversation patterns
            if any(keyword in para_lower for keyword in ['usuario:', 'cliente:', 'agente:', 'respuesta:']):
                patterns['conversation_patterns'].append({'line': i, 'content': paragraph})

        return {
            'structure_type': 'document',
            'patterns': patterns,
            'has_tables': len(raw_data['tables']) > 0,
            'table_structures': [{'rows': len(table), 'cols': len(table[0]) if table else 0} for table in raw_data['tables']]
        }

    def _analyze_pdf_structure(self, processed_file: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze PDF structure"""
        raw_data = processed_file['raw_data']

        # Similar to docx analysis but per page
        page_patterns = []

        for page_data in raw_data:
            content = page_data['content'].lower()
            page_pattern = {
                'page': page_data['page'],
                'has_tickets': any(keyword in content for keyword in ['ticket', 'caso', 'incidente']),
                'has_conversations': any(keyword in content for keyword in ['usuario:', 'cliente:', 'agente:']),
                'line_count': len(page_data['content'].split('\n'))
            }
            page_patterns.append(page_pattern)

        return {
            'structure_type': 'document',
            'page_patterns': page_patterns,
            'total_pages': len(page_patterns)
        }